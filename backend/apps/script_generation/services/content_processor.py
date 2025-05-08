# script_generation/services/content_processor.py
from django.conf import settings
import google.generativeai as genai
from PIL import Image
import pytesseract
from ...crawler.services.article_service import ArticleService  # Import từ crawler

class ContentProcessor:
    def __init__(self):
        genai.configure(api_key=settings.GEMINI_API_KEY)
        self.article_service = ArticleService()  # Sử dụng service từ crawler
    
    def get_articles_from_crawler(self, keyword=None, limit=5):
        """Gọi sang crawler service để lấy dữ liệu"""
        return self.article_service.get_articles(keyword, limit)

    def generate_simplified_explanation(self, content, audience="children", file_paths: list[str] = None):
        """
        Use Gemini to generate a simplified and easy-to-understand explanation tailored for a specific audience.

        Args:
            content (str): Scientific content to explain
            audience (str): Target audience (default: children)
            file_paths (list[str]): List of additional file paths (.txt, .png, .jpg, .pdf, .docx)

        Returns:
            str: Simplified explanation as text
        """
        from PIL import Image
        import pytesseract
        import fitz  # PyMuPDF
        import docx

        full_content = str(content) if content else ""
        file_content = ""

        if file_paths:
            for path in file_paths:
                try:
                    if not isinstance(path, str):
                        return f"[Error] Invalid file path: {path}"

                    if path.lower().endswith(('.png', '.jpg', '.jpeg')):
                        img = Image.open(path)
                        text = pytesseract.image_to_string(img, lang='eng')
                        file_content += f"\n\n[Text extracted from image {path}]:\n{text}"

                    elif path.lower().endswith('.txt'):
                        with open(path, 'r', encoding='utf-8') as f:
                            text = f.read()
                            file_content += f"\n\n[Text from .txt file {path}]:\n{text}"

                    elif path.lower().endswith('.pdf'):
                        doc = fitz.open(path)
                        text = "".join(page.get_text() for page in doc)
                        file_content += f"\n\n[Text from PDF file {path}]:\n{text}"

                    elif path.lower().endswith('.docx'):
                        document = docx.Document(path)
                        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                        file_content += f"\n\n[Text from DOCX file {path}]:\n{text}"

                    else:
                        file_content += f"\n\n[⚠️ Skipping unsupported file: {path}]"

                except Exception as e:
                    return f"[Error processing file {path}]: {str(e)}"

            if file_content:
                full_content = f"{full_content}\n\n{file_content}" if full_content else file_content

        if not str(full_content).strip():
            return "[Error] No input content to explain."

        prompt = f"""
    You are a science communication expert. Please explain the content below in a way that is suitable for an audience of **{audience}**.

    --- CONTENT TO EXPLAIN ---
    {full_content}

    --- REQUIREMENTS ---
    - Use clear, simple, and relatable language for {audience}
    - Avoid technical terms or abstract explanations
    - If possible, use examples, metaphors, or visual imagery to help {audience} understand more easily

    Present the explanation in the most vivid and accessible way possible.
    """

        try:
            model = genai.GenerativeModel('gemini-1.5-flash-latest')
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            return f"[AI Error] {str(e)}"

    def transform_to_story(content="", style="vivid", genre="adventure", file_paths: list[str] = None):
        """
        Convert scientific content into a story (from direct text input and file list)

        Args:
            content (str): Direct scientific content (default: empty)
            style (str): Storytelling style
            genre (str): Story genre
            file_paths (list[str]): List of file paths (image/text) with additional content

        Returns:
            str: Transformed story or error message
        """
        final_content = str(content) if content else ""
        file_content = ""

        if file_paths:
            for path in file_paths:
                try:
                    if not isinstance(path, str):
                        return f"[Error] Invalid file path: {path}"

                    if path.lower().endswith('.txt'):
                        with open(path, 'r', encoding='utf-8') as f:
                            text = f.read()
                            file_content += f"\n\n[Text from TXT file {path}]:\n{text}"

                    elif path.lower().endswith('.docx'):
                        import docx
                        doc = docx.Document(path)
                        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                        file_content += f"\n\n[Text from DOCX file {path}]:\n{text}"

                    elif path.lower().endswith('.pdf'):
                        import fitz
                        doc = fitz.open(path)
                        text = "".join(page.get_text() for page in doc)
                        file_content += f"\n\n[Text from PDF file {path}]:\n{text}"

                    elif path.lower().endswith('.doc'):
                        file_content += f"\n\n[⚠️ Skipping unsupported .doc file: {path}]"

                except Exception as e:
                    return f"[File read error {path}] {str(e)}"

            if file_content:
                final_content = f"{final_content}\n\n{file_content}" if final_content else file_content

        if not str(final_content).strip():
            return "[Error] No scientific content provided."

        prompt = f"""
You are a professional narrator and science storyteller. Turn the following scientific content into a vivid, structured, third-person story — told **only by a narrator** — in the **{style}** style and **{genre}** genre, suitable for **video and audio generation**.

--- ORIGINAL SCIENTIFIC CONTENT ---
{final_content}

--- STRUCTURE ---
Break the story into clearly numbered scenes like this:

### SCENE 1: [Short summary of the scene]
- **Narration (Third-person, no characters)**: Describe the scene in vivid, cinematic language. Focus on how the world changes, what is happening, and how science is involved. No dialogue or character names — just storytelling.
- **Visual Description (for still image generation)**: Describe what a single **illustrative image** for this scene should look like. Use precise, descriptive language suitable for generating a **single static image** (e.g., a high-resolution painting of a DNA strand twisting in space, a close-up photo of frost forming on a window).
- **Sound Effects / Music (optional)**: Suggest fitting background audio (e.g., gentle wind, deep rumble, uplifting orchestral music).

... Continue with SCENE 2, SCENE 3, etc.

### FINAL SECTION: THE SCIENCE BEHIND
- Under this heading, explain the scientific concepts presented in the story.
- Use simple metaphors or real-life comparisons to make them easy to understand.
- Keep it engaging and accessible to a general audience.

--- STYLE GUIDE ---
- Use only third-person narration. Do not include any characters or dialogues.
- Each scene should be 100–150 words (enough for ~30–60 seconds of audio).
- Language should be **cinematic, sensory, and suitable for narration**.
- Output in **English**.

Now begin the story.
"""

        try:
            model = genai.GenerativeModel('gemini-1.5-flash-latest')
            response = model.generate_content(prompt)
            return str(response.text)
        except Exception as e:
            return f"[AI Error] {str(e)}"

